from flask import Flask, request, render_template, redirect, url_for
import os
import pytesseract
from pdf2image import convert_from_path
from docx import Document
from PIL import Image
from werkzeug.utils import secure_filename

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads/'
app.config['ALLOWED_EXTENSIONS'] = {'pdf', 'docx', 'txt', 'jpg', 'jpeg', 'png'}

if not os.path.exists(app.config['UPLOAD_FOLDER']):
    os.makedirs(app.config['UPLOAD_FOLDER'])

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

def ocr_from_pdf(pdf_path):
    try:
        pages = convert_from_path(pdf_path, 300)
    except Exception as e:
        print(f"Error converting PDF to images: {e}")
        return ""

    text = ""
    for i, page in enumerate(pages):
        try:
            custom_config = r'--tessdata-dir "C:\Program Files\Tesseract-OCR\tessdata"'
            page_text = pytesseract.image_to_string(page, lang='fra', config=custom_config)
            text += page_text
        except Exception as e:
            print(f"Error performing OCR on page {i + 1}: {e}")
    return text

def extract_text_from_docx(docx_path):
    try:
        doc = Document(docx_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text
    except Exception as e:
        print(f"Error reading DOCX file: {e}")
        return ""

def extract_text_from_txt(txt_path):
    try:
        with open(txt_path, 'r', encoding='utf-8') as file:
            text = file.read()
        return text
    except Exception as e:
        print(f"Error reading TXT file: {e}")
        return ""

def ocr_from_image(image_path):
    try:
        image = Image.open(image_path)
        custom_config = r'--tessdata-dir "C:\Program Files\Tesseract-OCR\tessdata"'
        text = pytesseract.image_to_string(image, lang='fra', config=custom_config)
        return text
    except Exception as e:
        print(f"Error performing OCR on image: {e}")
        return ""

def save_to_word(text, docx_path):
    try:
        doc = Document()
        doc.add_heading('Extracted Text', 0)
        doc.add_paragraph(text)
        doc.save(docx_path)
    except Exception as e:
        print(f"Error saving Word document: {e}")

@app.route('/', methods=['GET', 'POST'])
def index():
    message = None
    if request.method == 'POST':
        if 'files' not in request.files:
            message = 'No file part'
            return render_template('index.html', message=message)
        
        files = request.files.getlist('files')
        output_folder = app.config['UPLOAD_FOLDER']
        
        for file in files:
            if file.filename == '':
                message = 'No selected file'
                return render_template('index.html', message=message)
            
            if not allowed_file(file.filename):
                message = f"File type not allowed: {file.filename}"
                return render_template('index.html', message=message)

            filename = secure_filename(file.filename)
            file_path = os.path.join(output_folder, filename)
            file.save(file_path)
            
            extension = filename.rsplit('.', 1)[1].lower()
            if extension == 'pdf':
                extracted_text = ocr_from_pdf(file_path)
            elif extension == 'docx':
                extracted_text = extract_text_from_docx(file_path)
            elif extension == 'txt':
                extracted_text = extract_text_from_txt(file_path)
            elif extension in {'jpg', 'jpeg', 'png'}:
                extracted_text = ocr_from_image(file_path)
            else:
                continue
            
            base_filename = os.path.splitext(filename)[0]
            docx_path = os.path.join(output_folder, f'{base_filename}_extracted_text.docx')
            save_to_word(extracted_text, docx_path)
        
        message = 'Files processed successfully!'
    
    return render_template('index.html', message=message)

if __name__ == '__main__':
    app.run(debug=True)
